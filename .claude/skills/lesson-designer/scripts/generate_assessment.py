#!/usr/bin/env python3
"""
Generate Dedicated Assessment Documents

Creates quizzes, tests, Socratic discussion guides, and performance tasks
with appropriate rubrics and grading criteria.

Requirements covered:
    - ASMT-02: Tool can generate dedicated assessment lessons (quizzes, tests)
    - ASMT-03: Tool can generate graded Socratic discussion guides
    - ASMT-04: Tool can generate performance tasks with rubrics
"""

import json
import os
import sys
from typing import Dict, Any, List, Optional
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Performance levels for rubrics
PERFORMANCE_LEVELS = [
    ("Advanced", 4, "Exceeds expectations"),
    ("Proficient", 3, "Meets expectations"),
    ("Developing", 2, "Approaching expectations"),
    ("Beginning", 1, "Below expectations")
]


def set_cell_shading(cell, color: str):
    """Set background shading for a table cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_assessment_header(doc: Document, title: str, assessment_type: str, total_points: int = None):
    """Add standard assessment header."""
    # Title
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Assessment type and date
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(f"{assessment_type.upper()}")
    run2.font.size = Pt(12)

    # Student info line
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cells = table.rows[0].cells
    cells[0].text = "Name: ____________________"
    cells[1].text = "Date: ________"
    cells[2].text = "Period: ____"
    if total_points:
        cells[3].text = f"Score: ____/{total_points}"
    else:
        cells[3].text = "Score: ____"

    for cell in cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)

    doc.add_paragraph()


def generate_multiple_choice_section(doc: Document, questions: List[Dict], start_num: int = 1) -> int:
    """Generate multiple choice questions section."""
    if not questions:
        return start_num

    doc.add_heading("Part 1: Multiple Choice", level=2)
    p = doc.add_paragraph("Circle the letter of the best answer.")
    p.runs[0].font.size = Pt(10)
    p.runs[0].italic = True

    q_num = start_num
    for q in questions:
        # Question
        q_para = doc.add_paragraph()
        run = q_para.add_run(f"{q_num}. {q['question']}")
        run.bold = True
        run.font.size = Pt(11)

        # Choices
        choices = q.get('choices', [])
        for i, choice in enumerate(choices):
            letter = chr(65 + i)  # A, B, C, D
            c_para = doc.add_paragraph(f"    {letter}. {choice}")
            c_para.paragraph_format.space_after = Pt(2)
            for run in c_para.runs:
                run.font.size = Pt(11)

        doc.add_paragraph()
        q_num += 1

    return q_num


def generate_short_answer_section(doc: Document, questions: List[Dict], start_num: int = 1) -> int:
    """Generate short answer questions section."""
    if not questions:
        return start_num

    doc.add_heading("Part 2: Short Answer", level=2)
    p = doc.add_paragraph("Write your answer in the space provided. Use complete sentences.")
    p.runs[0].font.size = Pt(10)
    p.runs[0].italic = True

    q_num = start_num
    for q in questions:
        points = q.get('points', 5)

        # Question with points
        q_para = doc.add_paragraph()
        run = q_para.add_run(f"{q_num}. {q['question']} ")
        run.bold = True
        run.font.size = Pt(11)

        pts_run = q_para.add_run(f"({points} points)")
        pts_run.font.size = Pt(10)
        pts_run.italic = True

        # Answer lines with double spacing
        lines = q.get('lines', 4)
        for _ in range(lines):
            answer_line = doc.add_paragraph("_" * 100)
            answer_line.paragraph_format.line_spacing = 2.0
            answer_line.paragraph_format.space_after = Pt(3)
            for run in answer_line.runs:
                run.font.size = Pt(11)

        doc.add_paragraph()
        q_num += 1

    return q_num


def generate_essay_section(doc: Document, questions: List[Dict], start_num: int = 1) -> int:
    """Generate essay questions section."""
    if not questions:
        return start_num

    doc.add_heading("Part 3: Extended Response", level=2)
    p = doc.add_paragraph("Write a well-organized response to the following prompt.")
    p.runs[0].font.size = Pt(10)
    p.runs[0].italic = True

    q_num = start_num
    for q in questions:
        points = q.get('points', 20)

        # Question with points
        q_para = doc.add_paragraph()
        run = q_para.add_run(f"{q_num}. {q['question']} ")
        run.bold = True
        run.font.size = Pt(11)

        pts_run = q_para.add_run(f"({points} points)")
        pts_run.font.size = Pt(10)
        pts_run.italic = True

        # Scoring criteria if provided
        if q.get('criteria'):
            crit_para = doc.add_paragraph()
            crit_run = crit_para.add_run("Your response will be evaluated on: ")
            crit_run.italic = True
            crit_run.font.size = Pt(10)
            for criterion in q['criteria']:
                crit_para.add_run(f"• {criterion}  ")

        # Answer space - use "Continue on back if needed"
        doc.add_paragraph()
        for _ in range(10):
            answer_line = doc.add_paragraph("_" * 100)
            answer_line.paragraph_format.line_spacing = 2.0
            answer_line.paragraph_format.space_after = Pt(2)

        continue_para = doc.add_paragraph("(Continue on back if needed)")
        continue_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        continue_para.runs[0].italic = True
        continue_para.runs[0].font.size = Pt(9)

        q_num += 1

    return q_num


def generate_quiz(assessment_data: Dict, output_path: str) -> bool:
    """Generate quiz document."""
    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Calculate total points
    total = 0
    for q in assessment_data.get('multiple_choice', []):
        total += q.get('points', 1)
    for q in assessment_data.get('short_answer', []):
        total += q.get('points', 5)

    # Header
    add_assessment_header(doc, assessment_data['title'], 'Quiz', total)

    # Sections
    q_num = 1
    q_num = generate_multiple_choice_section(doc, assessment_data.get('multiple_choice', []), q_num)
    q_num = generate_short_answer_section(doc, assessment_data.get('short_answer', []), q_num)

    doc.save(output_path)
    return True


def generate_test(assessment_data: Dict, output_path: str) -> bool:
    """Generate test document (more comprehensive than quiz)."""
    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Calculate total points
    total = 0
    for q in assessment_data.get('multiple_choice', []):
        total += q.get('points', 2)
    for q in assessment_data.get('short_answer', []):
        total += q.get('points', 5)
    for q in assessment_data.get('essay', []):
        total += q.get('points', 20)

    # Header
    add_assessment_header(doc, assessment_data['title'], 'Test', total)

    # Instructions
    inst_para = doc.add_paragraph()
    inst_para.add_run("Instructions: ").bold = True
    inst_para.add_run(assessment_data.get('instructions',
        'Read each question carefully. Show your work where applicable. '
        'You have the entire period to complete this test.'))
    inst_para.runs[1].font.size = Pt(10)

    doc.add_paragraph()

    # Sections
    q_num = 1
    q_num = generate_multiple_choice_section(doc, assessment_data.get('multiple_choice', []), q_num)
    q_num = generate_short_answer_section(doc, assessment_data.get('short_answer', []), q_num)
    q_num = generate_essay_section(doc, assessment_data.get('essay', []), q_num)

    doc.save(output_path)
    return True


def create_performance_rubric(doc: Document, task_name: str, criteria: List[Dict]):
    """Create analytical rubric table for performance task."""
    doc.add_heading("Scoring Rubric", level=2)

    # Create table: 1 header + criteria rows, 5 columns (criteria + 4 levels)
    num_rows = len(criteria) + 1
    table = doc.add_table(rows=num_rows, cols=5)
    table.style = 'Table Grid'

    # Header row
    header = table.rows[0].cells
    header[0].text = "Criteria"
    for i, (level_name, points, _) in enumerate(PERFORMANCE_LEVELS):
        header[i + 1].text = f"{level_name} ({points})"
        set_cell_shading(header[i + 1], "E8E8E8")
        for para in header[i + 1].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)

    set_cell_shading(header[0], "2D5A87")
    for para in header[0].paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Criteria rows
    for row_idx, criterion in enumerate(criteria):
        cells = table.rows[row_idx + 1].cells

        # Criterion name
        cells[0].text = criterion['name']
        set_cell_shading(cells[0], "F0F0F0")
        for para in cells[0].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)

        # Descriptors
        descriptors = criterion.get('descriptors', [''] * 4)
        for i, descriptor in enumerate(descriptors[:4]):
            cells[i + 1].text = descriptor
            for para in cells[i + 1].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    # Set column widths
    widths = [Inches(1.3), Inches(1.8), Inches(1.8), Inches(1.8), Inches(1.8)]
    for i, width in enumerate(widths):
        for cell in table.columns[i].cells:
            cell.width = width

    doc.add_paragraph()


def generate_performance_task(assessment_data: Dict, output_path: str) -> bool:
    """Generate performance task with rubric."""
    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Header
    add_assessment_header(doc, assessment_data['title'], 'Performance Task')

    # Task description
    doc.add_heading("Task Description", level=2)
    desc = doc.add_paragraph(assessment_data.get('description', ''))
    desc.runs[0].font.size = Pt(11)

    # Requirements
    if assessment_data.get('requirements'):
        doc.add_heading("Requirements", level=2)
        for req in assessment_data['requirements']:
            req_para = doc.add_paragraph(f"☐ {req}", style='List Bullet')
            for run in req_para.runs:
                run.font.size = Pt(11)

    # Rubric
    if assessment_data.get('criteria'):
        create_performance_rubric(doc, assessment_data['title'], assessment_data['criteria'])

    # Work space
    doc.add_heading("Your Work", level=2)
    workspace_note = doc.add_paragraph(
        "Use this space for planning. Attach your final product to this sheet.")
    workspace_note.runs[0].italic = True
    workspace_note.runs[0].font.size = Pt(10)

    for _ in range(15):
        line = doc.add_paragraph("_" * 100)
        line.paragraph_format.line_spacing = 2.0

    doc.save(output_path)
    return True


def generate_socratic_guide(assessment_data: Dict, output_path: str) -> bool:
    """Generate Socratic discussion guide with facilitation notes and rubric."""
    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    p = doc.add_paragraph()
    run = p.add_run(f"Socratic Seminar: {assessment_data['title']}")
    run.bold = True
    run.font.size = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Topic/Text
    doc.add_heading("Focus Text/Topic", level=2)
    topic = doc.add_paragraph(assessment_data.get('topic', assessment_data.get('description', '')))
    topic.runs[0].font.size = Pt(11)

    # Essential Questions
    doc.add_heading("Essential Questions", level=2)
    for i, q in enumerate(assessment_data.get('essential_questions', []), 1):
        q_para = doc.add_paragraph(f"{i}. {q}")
        q_para.runs[0].font.size = Pt(11)

    # Discussion Norms
    doc.add_heading("Discussion Norms", level=2)
    norms = assessment_data.get('norms', [
        "Listen actively and respectfully",
        "Build on others' ideas with 'I agree/disagree because...'",
        "Support claims with evidence from the text",
        "Ask clarifying questions",
        "Allow thinking time before responding"
    ])
    for norm in norms:
        norm_para = doc.add_paragraph(f"• {norm}")
        for run in norm_para.runs:
            run.font.size = Pt(10)

    # Participation Rubric
    socratic_criteria = assessment_data.get('criteria', [
        {
            'name': 'Participation',
            'descriptors': [
                'Actively contributes 4+ times with substantive comments',
                'Contributes 2-3 times with relevant ideas',
                'Contributes 1 time or with minimal depth',
                'Does not participate verbally'
            ]
        },
        {
            'name': 'Use of Evidence',
            'descriptors': [
                'Consistently cites specific textual evidence',
                'Sometimes cites evidence to support claims',
                'Rarely references text or evidence',
                'Makes claims without any evidence'
            ]
        },
        {
            'name': 'Engagement with Others',
            'descriptors': [
                'Builds on, challenges, or synthesizes peer ideas',
                'Responds to peer ideas occasionally',
                'Makes isolated comments unconnected to discussion',
                'Does not engage with peer contributions'
            ]
        }
    ])
    create_performance_rubric(doc, "Socratic Seminar", socratic_criteria)

    # Student Preparation Space
    doc.add_heading("Pre-Discussion Preparation", level=2)
    prep_note = doc.add_paragraph("Use this space to prepare your thoughts before the discussion.")
    prep_note.runs[0].italic = True

    doc.add_paragraph()
    doc.add_paragraph("My initial response to the essential question:")
    for _ in range(4):
        line = doc.add_paragraph("_" * 100)
        line.paragraph_format.line_spacing = 2.0

    doc.add_paragraph()
    doc.add_paragraph("Evidence from the text I want to reference:")
    for _ in range(4):
        line = doc.add_paragraph("_" * 100)
        line.paragraph_format.line_spacing = 2.0

    doc.add_paragraph()
    doc.add_paragraph("Questions I have for my peers:")
    for _ in range(3):
        line = doc.add_paragraph("_" * 100)
        line.paragraph_format.line_spacing = 2.0

    doc.save(output_path)
    return True


def generate_answer_key(assessment_data: Dict, output_path: str) -> bool:
    """Generate answer key for quiz/test."""
    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    p = doc.add_paragraph()
    run = p.add_run(f"ANSWER KEY: {assessment_data['title']}")
    run.bold = True
    run.font.size = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Multiple choice answers
    mc_questions = assessment_data.get('multiple_choice', [])
    if mc_questions:
        doc.add_heading("Multiple Choice Answers", level=2)
        for i, q in enumerate(mc_questions, 1):
            answer = q.get('answer', 'N/A')
            a_para = doc.add_paragraph(f"{i}. {answer}")
            for run in a_para.runs:
                run.font.size = Pt(11)

    # Short answer key points
    sa_questions = assessment_data.get('short_answer', [])
    if sa_questions:
        doc.add_heading("Short Answer Key Points", level=2)
        for i, q in enumerate(sa_questions, 1):
            doc.add_paragraph(f"{i}. {q['question']}")
            key_points = q.get('key_points', q.get('answer', 'See rubric'))
            if isinstance(key_points, list):
                for point in key_points:
                    kp = doc.add_paragraph(f"   • {point}")
                    for run in kp.runs:
                        run.font.size = Pt(10)
            else:
                kp = doc.add_paragraph(f"   {key_points}")
                for run in kp.runs:
                    run.font.size = Pt(10)

    # Essay rubric reference
    essay_questions = assessment_data.get('essay', [])
    if essay_questions:
        doc.add_heading("Extended Response Guidance", level=2)
        for i, q in enumerate(essay_questions, 1):
            doc.add_paragraph(f"{i}. {q['question']}")
            if q.get('key_points'):
                doc.add_paragraph("Key elements to look for:")
                for point in q['key_points']:
                    kp = doc.add_paragraph(f"   • {point}")
                    for run in kp.runs:
                        run.font.size = Pt(10)

    doc.save(output_path)
    return True


def generate_assessment(assessment_path: str, output_path: str = None) -> bool:
    """
    Main entry point: Generate assessment from JSON specification.

    Assessment JSON should have:
    - title: Assessment title
    - type: 'quiz', 'test', 'performance', or 'socratic'
    - questions/criteria based on type
    """
    # Load assessment data
    with open(assessment_path, 'r', encoding='utf-8') as f:
        assessment_data = json.load(f)

    assessment_type = assessment_data.get('type', 'quiz').lower()

    # Determine output path
    if output_path is None:
        base_dir = os.path.dirname(os.path.abspath(assessment_path))
        output_path = os.path.join(base_dir, f"{assessment_type}_assessment.docx")

    # Generate based on type
    if assessment_type == 'quiz':
        success = generate_quiz(assessment_data, output_path)
    elif assessment_type == 'test':
        success = generate_test(assessment_data, output_path)
    elif assessment_type == 'performance':
        success = generate_performance_task(assessment_data, output_path)
    elif assessment_type == 'socratic':
        success = generate_socratic_guide(assessment_data, output_path)
    else:
        print(f"Unknown assessment type: {assessment_type}", file=sys.stderr)
        return False

    if success:
        print(f"Generated {assessment_type}: {output_path}")

        # Generate answer key for quiz/test
        if assessment_type in ['quiz', 'test']:
            key_path = output_path.replace('.docx', '_answer_key.docx')
            generate_answer_key(assessment_data, key_path)
            print(f"Generated answer key: {key_path}")

    return success


def main():
    """CLI entry point."""
    if len(sys.argv) < 2:
        print("Usage: python generate_assessment.py <assessment.json> [output.docx]")
        print()
        print("Assessment types: quiz, test, performance, socratic")
        sys.exit(1)

    assessment_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) >= 3 else None

    if not os.path.exists(assessment_path):
        print(f"Error: Assessment file not found: {assessment_path}", file=sys.stderr)
        sys.exit(1)

    success = generate_assessment(assessment_path, output_path)
    sys.exit(0 if success else 1)


if __name__ == "__main__":
    main()
