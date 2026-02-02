#!/usr/bin/env python3
"""
Generate Compact Student Worksheet from Lesson Design JSON

Creates practical, printable worksheets (1-2 pages) using python-docx directly.
Uses tables for structured content and underscore answer lines.

Key features:
- Compact layout with Part 1, Part 2, Part 3 sections
- Tables for structured activities (ranking, definitions, etc.)
- Underscore answer lines for student responses
- 1-2 pages maximum for practical printing

Requirements covered:
    - MATL-01: Generate actual .docx files
    - MATL-03: Material type matches lesson type
    - ASMT-01: Each lesson includes assessment of its objective
"""

import json
import os
import sys
from typing import Dict, Any, List, Optional

from docx import Document
from docx.shared import Inches, Pt, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Default font for all documents
DEFAULT_FONT = 'Helvetica'

# Lines of answer space by activity type
WRITING_SPACE_CONFIG = {
    'retrieval': 2,           # Quick recall, less space needed
    'comprehension': 3,       # Explanations need moderate space
    'analysis': 5,            # Detailed responses need more space
    'knowledge_utilization': 6,  # Complex responses need most space
    'default': 3
}


def load_lesson_design(lesson_path: str) -> Dict[str, Any]:
    """Load lesson design from JSON file."""
    with open(lesson_path, 'r', encoding='utf-8') as f:
        return json.load(f)


def set_cell_shading(cell, color: str):
    """Set background shading for a table cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_header(doc: Document, lesson: Dict) -> None:
    """Add compact header with title and student info."""
    # Title
    title = lesson.get('title', 'Student Worksheet')
    grade = lesson.get('grade_level', '')

    p = doc.add_paragraph()
    run = p.add_run(f"{title}")
    run.bold = True
    run.font.size = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Name/Period/Date line - use a table for alignment
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    cells = table.rows[0].cells
    cells[0].text = "Name: _________________"
    cells[1].text = "Period: ____"
    cells[2].text = "Date: ____________"

    for cell in cells:
        cell.paragraphs[0].runs[0].font.size = Pt(10)

    # Small spacing
    doc.add_paragraph()


def add_part_header(doc: Document, part_num: int, title: str) -> None:
    """Add a Part header (e.g., 'Part 1: Do Now')."""
    p = doc.add_paragraph()
    run = p.add_run(f"Part {part_num}: {title}")
    run.bold = True
    run.font.size = Pt(11)
    p.space_after = Pt(6)


def add_instructions(doc: Document, text: str) -> None:
    """Add instruction text."""
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(10)


def add_image_placeholder(doc: Document, description: str, url: str = None) -> None:
    """Add a bordered placeholder box for an image with description and optional link."""
    # Create a single-cell table to act as a bordered box
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    cell = table.rows[0].cells[0]
    cell.width = Inches(5)

    # Add placeholder text
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run("[IMAGE: ")
    run.font.size = Pt(10)
    run.font.italic = True

    run2 = p.add_run(description[:100])
    run2.font.size = Pt(10)
    run2.font.italic = True

    run3 = p.add_run("]")
    run3.font.size = Pt(10)
    run3.font.italic = True

    # Add URL as hyperlink if provided
    if url:
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_hyperlink(p2, url, "View/Download Image")

    # Add some height to the cell for visual space
    p3 = cell.add_paragraph()
    p3.add_run("\n")

    # Add spacing after
    doc.add_paragraph()


def add_hyperlink(paragraph, url: str, text: str):
    """Add a hyperlink to a paragraph."""
    # Create the hyperlink element
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Style the hyperlink (blue, underlined)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '18')  # 9pt
    rPr.append(sz)

    new_run.append(rPr)
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_answer_lines(doc: Document, num_lines: int = 3, prefix: str = "") -> None:
    """Add underscore answer lines with double-spacing for student responses."""
    for i in range(num_lines):
        line_text = prefix if i == 0 and prefix else ""
        p = doc.add_paragraph(line_text + "_" * 80)
        p.paragraph_format.line_spacing = 2.0  # Double-spaced
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.size = Pt(11)  # Slightly larger for readability


def add_ranking_table(doc: Document, items: List[str], columns: List[str] = None) -> None:
    """Add a ranking/selection table with items."""
    if columns is None:
        columns = ["Item", "Rank", "Reason"]

    num_cols = len(columns)
    table = doc.add_table(rows=len(items) + 1, cols=num_cols)
    table.style = 'Table Grid'

    # Header row
    header_cells = table.rows[0].cells
    for i, col_name in enumerate(columns):
        header_cells[i].text = col_name
        header_cells[i].paragraphs[0].runs[0].bold = True
        header_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
        set_cell_shading(header_cells[i], "E0E0E0")

    # Data rows
    for row_idx, item in enumerate(items):
        cells = table.rows[row_idx + 1].cells
        cells[0].text = item
        cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        # Leave other cells empty for student input
        for cell in cells[1:]:
            cell.paragraphs[0].runs[0].font.size = Pt(10) if cell.paragraphs[0].runs else None

    doc.add_paragraph()  # Spacing after table


def add_definition_table(doc: Document, terms: List[Dict]) -> None:
    """Add a vocabulary/definition table."""
    table = doc.add_table(rows=len(terms) + 1, cols=2)
    table.style = 'Table Grid'

    # Header
    header = table.rows[0].cells
    header[0].text = "Term"
    header[1].text = "Definition"
    for cell in header:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        set_cell_shading(cell, "E0E0E0")

    # Terms
    for i, term in enumerate(terms):
        cells = table.rows[i + 1].cells
        cells[0].text = term.get('word', term.get('term', ''))
        cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        cells[0].paragraphs[0].runs[0].bold = True
        cells[1].text = term.get('definition', '')
        cells[1].paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()


def add_numbered_questions(doc: Document, questions: List[str], with_answer_space: bool = True, num_lines: int = 3) -> None:
    """Add numbered questions with answer space."""
    for i, question in enumerate(questions, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {question}")
        run.font.size = Pt(11)
        p.paragraph_format.line_spacing = 2.0  # Double-spaced
        p.paragraph_format.space_after = Pt(6)

        if with_answer_space:
            add_answer_lines(doc, num_lines)


def add_exit_ticket(doc: Document, questions: List[str]) -> None:
    """Add exit ticket section."""
    # Header with box effect using table
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    cell.text = "EXIT TICKET"
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(11)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading(cell, "E0E0E0")

    doc.add_paragraph()

    add_numbered_questions(doc, questions, with_answer_space=True, num_lines=4)


def generate_worksheet_from_lesson(lesson: Dict, output_path: str) -> bool:
    """
    Generate a compact worksheet from lesson design.

    Analyzes activities and creates appropriate sections with tables
    and answer lines. Targets 1-2 pages.
    """
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = DEFAULT_FONT
    style._element.rPr.rFonts.set(qn('w:eastAsia'), DEFAULT_FONT)

    # Set narrow margins for compact layout
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Add header
    add_header(doc, lesson)

    # Get activities
    activities = lesson.get('activities', [])

    # Generate parts based on activities
    part_num = 1

    for activity in activities:
        name = activity.get('name', f'Activity {part_num}')
        marzano_level = activity.get('marzano_level', 'comprehension')

        # Prefer student-facing content over teacher instructions
        student_directions = activity.get('student_directions', '')
        student_questions = activity.get('student_questions', [])
        discussion_questions = activity.get('discussion_questions', [])
        visual_description = activity.get('visual_description', '')
        recommended_image = activity.get('recommended_image', {})

        # Determine answer line count based on cognitive complexity
        answer_line_count = WRITING_SPACE_CONFIG.get(marzano_level, WRITING_SPACE_CONFIG['default'])

        # Skip exit ticket activities - we'll add that separately
        if 'exit' in name.lower() and 'ticket' in name.lower():
            continue

        # Add part header
        add_part_header(doc, part_num, name)

        # Add visual placeholder if specified (with URL if available from image search)
        if visual_description or recommended_image:
            img_desc = recommended_image.get('description', visual_description) or visual_description
            img_url = recommended_image.get('url', '')
            add_image_placeholder(doc, img_desc, img_url)

        # Add student directions (the actual task prompt students see)
        if student_directions:
            add_instructions(doc, student_directions)

        # Determine content type based on activity characteristics
        # Check for ranking/comparison activities
        if any(word in name.lower() for word in ['rank', 'order', 'compare', 'sort']):
            items_to_rank = activity.get('items', ['Item A', 'Item B', 'Item C', 'Item D'])
            if isinstance(items_to_rank, list) and len(items_to_rank) > 0:
                add_ranking_table(doc, items_to_rank[:5])
            add_instructions(doc, "Explain your #1 choice:")
            add_answer_lines(doc, answer_line_count)

        # Check for definition/vocabulary activities
        elif any(word in name.lower() for word in ['definition', 'vocabulary', 'terms', 'key words']):
            vocab = lesson.get('vocabulary', [])
            if vocab:
                add_definition_table(doc, vocab[:6])
            else:
                add_answer_lines(doc, answer_line_count)

        # Add student questions if present (primary content for worksheets)
        elif student_questions:
            add_numbered_questions(doc, student_questions, with_answer_space=True, num_lines=answer_line_count)

        # Add discussion questions if present
        elif discussion_questions:
            add_numbered_questions(doc, discussion_questions, with_answer_space=True, num_lines=answer_line_count)

        # Check for analysis/application activities
        elif marzano_level in ['analysis', 'knowledge_utilization']:
            student_output = activity.get('student_output', '')
            if student_output:
                add_instructions(doc, f"Your task: {student_output}")
            add_answer_lines(doc, answer_line_count)

        # Default: simple answer space
        else:
            add_answer_lines(doc, answer_line_count)

        part_num += 1

    # Add vocabulary section if present and not already added
    vocab = lesson.get('vocabulary', [])
    if vocab and part_num <= 3:  # Only add if we haven't had many parts
        add_part_header(doc, part_num, "Key Terms")
        add_definition_table(doc, vocab[:6])
        part_num += 1

    # Add exit ticket
    assessment = lesson.get('assessment', {})
    exit_questions = assessment.get('questions', [])

    if not exit_questions:
        # Default exit ticket questions
        objective = lesson.get('objective', 'the main concept')
        exit_questions = [
            f"What is the most important thing you learned today?",
            f"What is one question you still have?"
        ]

    add_exit_ticket(doc, exit_questions[:3])  # Max 3 questions

    # Save document
    doc.save(output_path)
    return True


def validate_output(output_path: str) -> tuple:
    """Verify generated file is valid and compact."""
    errors = []
    warnings = []

    if not os.path.exists(output_path):
        errors.append(f"File not found: {output_path}")
        return False, errors, warnings

    try:
        doc = Document(output_path)
    except Exception as e:
        errors.append(f"Invalid Word document: {e}")
        return False, errors, warnings

    # Check paragraph count (compact should be under 50)
    para_count = len([p for p in doc.paragraphs if p.text.strip()])
    if para_count > 60:
        warnings.append(f"Document may be too long ({para_count} paragraphs)")

    # Check for tables (should have at least one for compact format)
    if len(doc.tables) == 0:
        warnings.append("No tables found - consider using tables for structured content")

    # Check for double-spacing in answer paragraphs
    answer_paragraphs = [p for p in doc.paragraphs if '_' in p.text]
    if answer_paragraphs:
        # Check if at least some answer lines have double-spacing
        double_spaced_count = sum(1 for p in answer_paragraphs
                                  if p.paragraph_format.line_spacing and p.paragraph_format.line_spacing >= 1.5)
        if double_spaced_count < len(answer_paragraphs) * 0.5:
            warnings.append(f"Only {double_spaced_count}/{len(answer_paragraphs)} answer lines have adequate spacing (expected >= 1.5)")
    else:
        warnings.append("No answer lines found - worksheet may lack adequate answer space")

    return True, errors, warnings


def generate_worksheet(lesson_path: str, template_path: str = None, output_path: str = None) -> bool:
    """
    Generate student worksheet from lesson design.

    Note: template_path is ignored - we build documents programmatically for compactness.
    """
    # Load lesson
    lesson = load_lesson_design(lesson_path)

    # Determine output path
    if output_path is None:
        lesson_dir = os.path.dirname(os.path.abspath(lesson_path))
        output_path = os.path.join(lesson_dir, '06_worksheet.docx')

    # Generate worksheet
    print(f"Lesson type: {lesson.get('lesson_type', 'unknown')}")
    print(f"Generating compact worksheet...")

    success = generate_worksheet_from_lesson(lesson, output_path)

    if success:
        is_valid, errors, warnings = validate_output(output_path)

        if errors:
            print("\nErrors:", file=sys.stderr)
            for e in errors:
                print(f"  - {e}", file=sys.stderr)

        if warnings:
            print("\nWarnings:")
            for w in warnings:
                print(f"  - {w}")

        print(f"\nWorksheet generated successfully: {output_path}")
        return True

    return False


def main():
    """CLI entry point."""
    if len(sys.argv) < 2:
        print("Usage: python generate_worksheet.py <lesson.json> [template.docx] [output.docx]")
        print()
        print("Note: template.docx is ignored - documents are built programmatically for compactness.")
        sys.exit(1)

    lesson_path = sys.argv[1]

    if not os.path.exists(lesson_path):
        print(f"Error: Lesson file not found: {lesson_path}", file=sys.stderr)
        sys.exit(1)

    # Get output path (ignore template - we don't use it)
    output_path = sys.argv[3] if len(sys.argv) >= 4 else None
    if output_path is None and len(sys.argv) >= 3:
        # Check if arg 2 looks like an output path
        if sys.argv[2].endswith('.docx') and 'template' not in sys.argv[2].lower():
            output_path = sys.argv[2]

    success = generate_worksheet(lesson_path, None, output_path)
    sys.exit(0 if success else 1)


if __name__ == "__main__":
    main()
