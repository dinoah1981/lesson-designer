#!/usr/bin/env python3
"""
Generate Sequence-Level Assessments

Creates comprehensive assessments spanning multiple lessons in a sequence,
including cumulative tests, performance tasks, and portfolio reviews.

This module provides backward design assessment generation - assessments
measure end-of-sequence competency mastery across 2-4 week units.
"""

import json
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional, List, Dict, Any

from sequence_manager import get_sequence_metadata, get_lesson_directory
from generate_assessment import (
    create_document_with_font,
    add_assessment_header,
    create_performance_rubric,
    set_cell_shading,
    PERFORMANCE_LEVELS
)
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


@dataclass
class SequenceAssessmentConfig:
    """Configuration for sequence assessment generation."""
    assessment_type: str  # "cumulative_test", "performance_task", "portfolio_review"
    title: str
    time_limit: int = 60  # minutes
    include_lessons: Optional[List[int]] = None  # None = all lessons
    emphasis_competencies: Optional[List[str]] = None  # Competency IDs to emphasize


def generate_sequence_assessment(
    sequence_id: str,
    config: SequenceAssessmentConfig
) -> Dict[str, Any]:
    """
    Generate a sequence-level assessment covering multiple lessons.

    Args:
        sequence_id: The unique sequence identifier
        config: Assessment configuration specifying type and parameters

    Returns:
        dict: Assessment data structure matching generate_assessment.py schema

    Example:
        >>> config = SequenceAssessmentConfig(
        ...     assessment_type="cumulative_test",
        ...     title="Unit 1: Primary Sources Assessment",
        ...     time_limit=50
        ... )
        >>> assessment = generate_sequence_assessment(seq_id, config)
    """
    # Load sequence metadata
    metadata = get_sequence_metadata(sequence_id)

    # Load lesson summaries
    summaries = []
    lessons_to_include = config.include_lessons or list(range(1, metadata['total_lessons'] + 1))

    for lesson_num in lessons_to_include:
        lesson_dir = get_lesson_directory(sequence_id, lesson_num)
        summary_path = lesson_dir / 'lesson_summary.json'

        if summary_path.exists():
            with open(summary_path, 'r', encoding='utf-8') as f:
                summaries.append(json.load(f))

    # Build assessment based on type
    if config.assessment_type == "cumulative_test":
        assessment = _build_cumulative_test(metadata, summaries, config)
    elif config.assessment_type == "performance_task":
        assessment = _build_performance_task(metadata, summaries, config)
    elif config.assessment_type == "portfolio_review":
        assessment = _build_portfolio_review(metadata, summaries, config)
    else:
        raise ValueError(f"Unknown assessment type: {config.assessment_type}")

    return assessment


def _build_cumulative_test(
    metadata: Dict,
    summaries: List[Dict],
    config: SequenceAssessmentConfig
) -> Dict[str, Any]:
    """
    Build a cumulative test with questions at varied Marzano levels.

    Includes:
    - Vocabulary retrieval questions
    - Comprehension questions (explain concepts)
    - Analysis questions (compare/evaluate)
    - Knowledge utilization (apply to new scenarios)

    Args:
        metadata: Sequence metadata dict
        summaries: List of lesson summary dicts
        config: Assessment configuration

    Returns:
        dict: Test assessment data
    """
    # Collect vocabulary from progression
    vocabulary_terms = []
    vocab_progression = metadata.get('vocabulary_progression', {})
    for term, lessons in vocab_progression.items():
        vocabulary_terms.append(term)

    # Collect competencies
    competencies = metadata.get('competencies', [])
    emphasis_comp_ids = config.emphasis_competencies or [c['id'] for c in competencies]

    # Build multiple choice section (retrieval + comprehension)
    multiple_choice = []

    # Vocabulary retrieval questions (1 point each)
    for i, term in enumerate(vocabulary_terms[:5]):  # Max 5 vocab questions
        multiple_choice.append({
            'question': f'Which definition best describes "{term}"?',
            'choices': [
                f'{term} means [correct definition here]',
                f'{term} means [incorrect option 1]',
                f'{term} means [incorrect option 2]',
                f'{term} means [incorrect option 3]'
            ],
            'answer': 'A',
            'points': 1,
            'marzano_level': 'retrieval'
        })

    # Comprehension questions (2 points each)
    for i, comp in enumerate([c for c in competencies if c['id'] in emphasis_comp_ids][:3]):
        multiple_choice.append({
            'question': f'What is the main purpose of {comp["statement"].lower()}?',
            'choices': [
                'To apply critical thinking skills [correct]',
                'To memorize facts',
                'To complete worksheets',
                'To prepare for the next grade'
            ],
            'answer': 'A',
            'points': 2,
            'marzano_level': 'comprehension'
        })

    # Build short answer section (analysis)
    short_answer = []

    for i, comp in enumerate([c for c in competencies if c['id'] in emphasis_comp_ids][:2]):
        short_answer.append({
            'question': f'Compare two different approaches to {comp["statement"].lower()}. '
                       f'Which approach is more effective and why?',
            'points': 8,
            'lines': 6,
            'key_points': [
                'Identifies two distinct approaches',
                'Provides comparison of strengths/weaknesses',
                'Justifies selection with reasoning'
            ],
            'marzano_level': 'analysis'
        })

    # Build essay section (knowledge utilization)
    essay = []

    # Synthesis question requiring integration of all competencies
    comp_statements = [c['statement'] for c in competencies if c['id'] in emphasis_comp_ids]
    essay.append({
        'question': f'Given a new scenario [describe scenario here], apply what you have learned '
                   f'about {", ".join(comp_statements[:2])} to solve this problem. '
                   f'Explain your reasoning and justify your approach.',
        'points': 20,
        'criteria': [
            'Application of multiple competencies',
            'Clear reasoning and justification',
            'Use of vocabulary from the unit',
            'Well-organized response'
        ],
        'key_points': [
            'Correctly identifies which competencies apply',
            'Integrates knowledge from multiple lessons',
            'Uses unit vocabulary appropriately',
            'Provides step-by-step justification'
        ],
        'marzano_level': 'knowledge_utilization'
    })

    return {
        'title': config.title,
        'type': 'test',
        'instructions': f'You have {config.time_limit} minutes to complete this test. '
                       f'Read each question carefully and show your work.',
        'multiple_choice': multiple_choice,
        'short_answer': short_answer,
        'essay': essay,
        'metadata': {
            'sequence_id': metadata['sequence_id'],
            'lessons_covered': config.include_lessons or list(range(1, metadata['total_lessons'] + 1)),
            'time_limit': config.time_limit
        }
    }


def _build_performance_task(
    metadata: Dict,
    summaries: List[Dict],
    config: SequenceAssessmentConfig
) -> Dict[str, Any]:
    """
    Build a performance task requiring competency integration.

    Creates authentic task with rubric incorporating criteria from
    each competency. Uses backward design aligned to proficiency targets.

    Args:
        metadata: Sequence metadata dict
        summaries: List of lesson summary dicts
        config: Assessment configuration

    Returns:
        dict: Performance task assessment data
    """
    competencies = metadata.get('competencies', [])
    emphasis_comp_ids = config.emphasis_competencies or [c['id'] for c in competencies]
    emphasis_comps = [c for c in competencies if c['id'] in emphasis_comp_ids]

    # Build integrated task description
    comp_statements = [c['statement'] for c in emphasis_comps]
    task_description = (
        f"You will demonstrate mastery of this unit's competencies by completing "
        f"an authentic task that requires you to {comp_statements[0].lower()}"
    )

    if len(comp_statements) > 1:
        task_description += f" and {comp_statements[1].lower()}"

    task_description += (
        f". Your work will be evaluated on how well you integrate "
        f"knowledge and skills from all {len(summaries)} lessons in this sequence."
    )

    # Build requirements checklist
    requirements = []
    for comp in emphasis_comps:
        requirements.append(f"Demonstrate {comp['statement'].lower()}")

    # Collect vocabulary to include
    vocab_terms = list(metadata.get('vocabulary_progression', {}).keys())
    if vocab_terms:
        requirements.append(f"Use at least 3 vocabulary terms from the unit: {', '.join(vocab_terms[:5])}")

    requirements.append("Provide evidence and justification for your decisions")
    requirements.append("Present your work in a clear, organized format")

    # Build rubric criteria (one per competency + communication)
    criteria = []

    for comp in emphasis_comps:
        criteria.append({
            'name': comp['statement'][:30],  # Truncate long statements
            'descriptors': [
                f'Demonstrates sophisticated application of {comp["statement"].lower()[:20]}... with nuanced understanding',
                f'Demonstrates competent application of {comp["statement"].lower()[:20]}... with clear understanding',
                f'Demonstrates developing application of {comp["statement"].lower()[:20]}... with partial understanding',
                f'Demonstrates minimal application of {comp["statement"].lower()[:20]}... with limited understanding'
            ]
        })

    # Add communication criterion
    criteria.append({
        'name': 'Communication & Organization',
        'descriptors': [
            'Exceptionally clear, well-organized presentation with proper vocabulary usage',
            'Clear, organized presentation with appropriate vocabulary usage',
            'Somewhat organized presentation with some vocabulary usage',
            'Unclear or disorganized presentation with limited vocabulary'
        ]
    })

    return {
        'title': config.title,
        'type': 'performance',
        'description': task_description,
        'requirements': requirements,
        'criteria': criteria,
        'metadata': {
            'sequence_id': metadata['sequence_id'],
            'lessons_covered': config.include_lessons or list(range(1, metadata['total_lessons'] + 1)),
            'time_limit': config.time_limit
        }
    }


def _build_portfolio_review(
    metadata: Dict,
    summaries: List[Dict],
    config: SequenceAssessmentConfig
) -> Dict[str, Any]:
    """
    Build a portfolio review with reflection prompts and self-assessment.

    Args:
        metadata: Sequence metadata dict
        summaries: List of lesson summary dicts
        config: Assessment configuration

    Returns:
        dict: Portfolio review assessment data
    """
    competencies = metadata.get('competencies', [])
    lessons_to_include = config.include_lessons or list(range(1, metadata['total_lessons'] + 1))

    # Build reflection prompts (one per lesson)
    reflection_prompts = []

    for lesson_num in lessons_to_include:
        reflection_prompts.append({
            'lesson': lesson_num,
            'prompt': f'What was the most important thing you learned in Lesson {lesson_num}? '
                     f'How has your understanding changed since the beginning of the unit?',
            'lines': 6
        })

    # Add overall reflection
    reflection_prompts.append({
        'lesson': 'Overall',
        'prompt': f'Looking at all {len(lessons_to_include)} lessons, how have you grown in your ability to '
                 f'{competencies[0]["statement"].lower() if competencies else "apply what you learned"}? '
                 f'What evidence from your work demonstrates this growth?',
        'lines': 8
    })

    # Build self-assessment rubric
    self_assessment_criteria = []

    for comp in competencies[:3]:  # Max 3 competencies for self-assessment
        self_assessment_criteria.append({
            'name': comp['statement'][:40],
            'descriptors': [
                f'I can independently {comp["statement"].lower()[:30]}... in new situations',
                f'I can {comp["statement"].lower()[:30]}... with occasional help',
                f'I am developing the ability to {comp["statement"].lower()[:30]}...',
                f'I am just beginning to {comp["statement"].lower()[:30]}...'
            ]
        })

    return {
        'title': config.title,
        'type': 'portfolio',
        'description': f'Reflect on your learning across {len(lessons_to_include)} lessons and '
                      f'assess your own growth in the unit competencies.',
        'reflection_prompts': reflection_prompts,
        'self_assessment': self_assessment_criteria,
        'metadata': {
            'sequence_id': metadata['sequence_id'],
            'lessons_covered': lessons_to_include,
            'time_limit': config.time_limit
        }
    }


def _create_assessment_docx(
    assessment: Dict[str, Any],
    output_dir: Path
) -> Path:
    """
    Generate Word document for the assessment.

    Args:
        assessment: Assessment data dict
        output_dir: Directory to save the document

    Returns:
        Path: Absolute path to generated .docx file
    """
    doc = create_document_with_font()

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    assessment_type = assessment['type']

    if assessment_type == 'test':
        return _create_test_docx(assessment, output_dir)
    elif assessment_type == 'performance':
        return _create_performance_docx(assessment, output_dir)
    elif assessment_type == 'portfolio':
        return _create_portfolio_docx(assessment, output_dir)
    else:
        raise ValueError(f"Unknown assessment type for DOCX: {assessment_type}")


def _create_test_docx(assessment: Dict, output_dir: Path) -> Path:
    """Generate test document using existing generate_assessment functions."""
    from generate_assessment import generate_test

    # Save assessment JSON temporarily
    temp_json = output_dir / 'temp_assessment.json'
    with open(temp_json, 'w', encoding='utf-8') as f:
        json.dump(assessment, f, indent=2)

    # Generate using existing function
    output_path = output_dir / f"{assessment['title'].replace(' ', '_')}.docx"
    generate_test(assessment, str(output_path))

    # Clean up temp file
    temp_json.unlink()

    # Generate answer key
    from generate_assessment import generate_answer_key
    key_path = output_dir / f"{assessment['title'].replace(' ', '_')}_answer_key.docx"
    generate_answer_key(assessment, str(key_path))

    return output_path


def _create_performance_docx(assessment: Dict, output_dir: Path) -> Path:
    """Generate performance task document."""
    from generate_assessment import generate_performance_task

    output_path = output_dir / f"{assessment['title'].replace(' ', '_')}.docx"
    generate_performance_task(assessment, str(output_path))

    return output_path


def _create_portfolio_docx(assessment: Dict, output_dir: Path) -> Path:
    """Generate portfolio review document."""
    doc = create_document_with_font()

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Header
    add_assessment_header(doc, assessment['title'], 'Portfolio Review')

    # Description
    doc.add_heading("Purpose", level=2)
    desc = doc.add_paragraph(assessment['description'])
    desc.runs[0].font.size = Pt(11)

    # Reflection prompts
    doc.add_heading("Lesson Reflections", level=2)
    for prompt_data in assessment.get('reflection_prompts', []):
        lesson_label = f"Lesson {prompt_data['lesson']}" if isinstance(prompt_data['lesson'], int) else prompt_data['lesson']

        p = doc.add_paragraph()
        run = p.add_run(f"{lesson_label}: ")
        run.bold = True
        run.font.size = Pt(11)
        p.add_run(prompt_data['prompt'])

        # Answer lines
        lines = prompt_data.get('lines', 5)
        for _ in range(lines):
            line = doc.add_paragraph("_" * 80)
            line.paragraph_format.line_spacing = 2.0

        doc.add_paragraph()

    # Self-assessment rubric
    if assessment.get('self_assessment'):
        doc.add_heading("Self-Assessment Rubric", level=2)
        doc.add_paragraph("Rate yourself on each competency by circling the description that best matches your current ability:")

        create_performance_rubric(doc, "Self-Assessment", assessment['self_assessment'])

    output_path = output_dir / f"{assessment['title'].replace(' ', '_')}.docx"
    doc.save(str(output_path))

    return output_path


if __name__ == '__main__':
    print("Testing generate_sequence_assessment.py...")
    print("Import successful!")
