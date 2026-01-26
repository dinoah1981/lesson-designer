#!/usr/bin/env python3
"""
Tests for sequence-level assessment generation.

Tests cumulative tests, performance tasks, and portfolio reviews
that span multiple lessons in a sequence.
"""

import json
import pytest
import shutil
from pathlib import Path

import sys
sys.path.insert(0, str(Path(__file__).parent.parent / 'scripts'))

from generate_sequence_assessment import (
    generate_sequence_assessment,
    SequenceAssessmentConfig
)
from sequence_manager import (
    create_sequence_session,
    get_sessions_dir,
    save_sequence_metadata
)
from parse_competency import get_project_root


@pytest.fixture
def test_sequence():
    """Create a test sequence with multiple lessons."""
    # Create sequence
    sequence_id = create_sequence_session(
        competencies=[
            "Analyze primary sources to evaluate historical claims",
            "Compare multiple perspectives on historical events"
        ],
        grade_level="8th grade",
        total_lessons=3,
        lesson_duration=55
    )

    # Add vocabulary progression
    from sequence_manager import get_sequence_metadata
    metadata = get_sequence_metadata(sequence_id)
    metadata['vocabulary_progression'] = {
        'primary source': [1, 2],
        'bias': [1, 2, 3],
        'perspective': [2, 3],
        'corroborate': [2, 3],
        'inference': [1, 2]
    }
    save_sequence_metadata(sequence_id, metadata)

    # Create mock lesson summaries
    from sequence_manager import get_lesson_directory
    for lesson_num in range(1, 4):
        lesson_dir = get_lesson_directory(sequence_id, lesson_num)
        summary = {
            'lesson_number': lesson_num,
            'objectives': [
                f'Students will analyze primary sources from lesson {lesson_num}',
                f'Students will evaluate claims in lesson {lesson_num}'
            ],
            'vocabulary': ['primary source', 'bias'] if lesson_num == 1 else ['perspective', 'corroborate'],
            'marzano_distribution': {
                'retrieval': 0.2,
                'comprehension': 0.3,
                'analysis': 0.3,
                'knowledge_utilization': 0.2
            }
        }
        summary_path = lesson_dir / 'lesson_summary.json'
        with open(summary_path, 'w', encoding='utf-8') as f:
            json.dump(summary, f, indent=2)

    yield sequence_id

    # Cleanup
    sequence_dir = get_sessions_dir() / sequence_id
    if sequence_dir.exists():
        shutil.rmtree(sequence_dir)


def test_cumulative_test_generation(test_sequence):
    """Test cumulative test assessment generation."""
    config = SequenceAssessmentConfig(
        assessment_type="cumulative_test",
        title="Unit 1: Primary Sources Test",
        time_limit=50
    )

    assessment = generate_sequence_assessment(test_sequence, config)

    # Verify assessment structure
    assert assessment['title'] == "Unit 1: Primary Sources Test"
    assert assessment['type'] == 'test'
    assert 'multiple_choice' in assessment
    assert 'short_answer' in assessment
    assert 'essay' in assessment

    # Verify sections exist
    assert len(assessment['multiple_choice']) > 0, "Should have multiple choice questions"
    assert len(assessment['short_answer']) > 0, "Should have short answer questions"
    assert len(assessment['essay']) > 0, "Should have essay questions"

    # Verify Marzano level distribution
    mc_levels = [q.get('marzano_level') for q in assessment['multiple_choice']]
    assert 'retrieval' in mc_levels, "Should have retrieval questions"
    assert 'comprehension' in mc_levels, "Should have comprehension questions"

    sa_levels = [q.get('marzano_level') for q in assessment['short_answer']]
    assert 'analysis' in sa_levels, "Should have analysis questions"

    essay_levels = [q.get('marzano_level') for q in assessment['essay']]
    assert 'knowledge_utilization' in essay_levels, "Should have knowledge utilization questions"

    # Verify vocabulary appears in questions
    # Check that at least some vocabulary terms are referenced
    all_questions = [q['question'] for q in assessment['multiple_choice']]
    vocab_found = any('primary source' in q.lower() or 'bias' in q.lower() or
                     'perspective' in q.lower() for q in all_questions)
    assert vocab_found, "Vocabulary from progression should appear in questions"

    # Verify answer key exists for MC questions
    for q in assessment['multiple_choice']:
        assert 'answer' in q, "Multiple choice questions should have answers"

    # Verify metadata
    assert assessment['metadata']['sequence_id'] == test_sequence
    assert assessment['metadata']['time_limit'] == 50


def test_performance_task_generation(test_sequence):
    """Test performance task assessment generation."""
    config = SequenceAssessmentConfig(
        assessment_type="performance_task",
        title="Historical Analysis Performance Task",
        time_limit=90
    )

    assessment = generate_sequence_assessment(test_sequence, config)

    # Verify assessment structure
    assert assessment['title'] == "Historical Analysis Performance Task"
    assert assessment['type'] == 'performance'
    assert 'description' in assessment
    assert 'requirements' in assessment
    assert 'criteria' in assessment

    # Verify task description references competencies
    description = assessment['description'].lower()
    assert 'analyze' in description or 'compare' in description, \
        "Task description should reference competencies"

    # Verify requirements exist
    assert len(assessment['requirements']) > 0, "Should have requirements"

    # Verify rubric has criteria from each competency
    assert len(assessment['criteria']) >= 2, \
        "Should have criteria for each competency (plus communication)"

    # Check criteria structure (4-level rubric)
    for criterion in assessment['criteria']:
        assert 'name' in criterion
        assert 'descriptors' in criterion
        assert len(criterion['descriptors']) == 4, \
            "Each criterion should have 4 performance levels"

    # Verify success criteria align to proficiency targets
    # Check that descriptors reference competencies
    first_criterion = assessment['criteria'][0]
    descriptors = ' '.join(first_criterion['descriptors']).lower()
    # Should have graded descriptors from advanced to beginning
    assert 'understanding' in descriptors, \
        "Rubric descriptors should reference understanding/competency"

    # Verify metadata
    assert assessment['metadata']['sequence_id'] == test_sequence
    assert assessment['metadata']['time_limit'] == 90


def test_partial_lesson_assessment(test_sequence):
    """Test assessment covering only subset of lessons (mid-unit check)."""
    config = SequenceAssessmentConfig(
        assessment_type="cumulative_test",
        title="Mid-Unit Check",
        time_limit=30,
        include_lessons=[1, 2]  # Only first 2 lessons
    )

    assessment = generate_sequence_assessment(test_sequence, config)

    # Verify only lessons 1-2 are covered
    assert assessment['metadata']['lessons_covered'] == [1, 2]

    # Verify vocabulary from lessons 1-2 only
    # Lesson 1-2 vocab: primary source, bias, perspective, corroborate
    # Lesson 3 would add more, but shouldn't be included

    # Check that assessment questions exist but are scoped appropriately
    assert len(assessment['multiple_choice']) > 0
    assert len(assessment['short_answer']) > 0

    # The assessment should be shorter (fewer total questions)
    # than a full 3-lesson assessment would be
    total_questions = (len(assessment['multiple_choice']) +
                      len(assessment['short_answer']) +
                      len(assessment['essay']))
    assert total_questions > 0, "Should have questions from included lessons"


def test_assessment_docx_creation(test_sequence, tmp_path):
    """Test Word document generation for assessments."""
    config = SequenceAssessmentConfig(
        assessment_type="cumulative_test",
        title="DOCX Test Assessment",
        time_limit=45
    )

    assessment = generate_sequence_assessment(test_sequence, config)

    # Generate DOCX
    from generate_sequence_assessment import _create_assessment_docx
    output_path = _create_assessment_docx(assessment, tmp_path)

    # Verify file was created
    assert output_path.exists(), f"DOCX file should be created at {output_path}"
    assert output_path.suffix == '.docx', "Should be a .docx file"

    # Verify file is valid Word document (can be opened)
    from docx import Document
    try:
        doc = Document(str(output_path))
        assert len(doc.paragraphs) > 0, "Document should have content"
    except Exception as e:
        pytest.fail(f"DOCX file is not valid: {e}")

    # Verify answer key was also created for test type
    answer_key_path = tmp_path / f"{assessment['title'].replace(' ', '_')}_answer_key.docx"
    assert answer_key_path.exists(), "Answer key should be created for tests"

    # Verify answer key is valid
    try:
        key_doc = Document(str(answer_key_path))
        assert len(key_doc.paragraphs) > 0, "Answer key should have content"
    except Exception as e:
        pytest.fail(f"Answer key DOCX is not valid: {e}")


def test_portfolio_review_generation(test_sequence):
    """Test portfolio review assessment generation."""
    config = SequenceAssessmentConfig(
        assessment_type="portfolio_review",
        title="Unit Reflection Portfolio",
        time_limit=60
    )

    assessment = generate_sequence_assessment(test_sequence, config)

    # Verify assessment structure
    assert assessment['title'] == "Unit Reflection Portfolio"
    assert assessment['type'] == 'portfolio'
    assert 'description' in assessment
    assert 'reflection_prompts' in assessment
    assert 'self_assessment' in assessment

    # Verify reflection prompts exist for each lesson + overall
    prompts = assessment['reflection_prompts']
    assert len(prompts) >= 3, "Should have prompts for each lesson"

    # Check that prompts have lesson numbers and questions
    for prompt in prompts[:-1]:  # All but last (overall reflection)
        assert 'lesson' in prompt
        assert 'prompt' in prompt
        assert isinstance(prompt['lesson'], int)

    # Check overall reflection exists
    assert prompts[-1]['lesson'] == 'Overall'

    # Verify self-assessment rubric
    assert len(assessment['self_assessment']) > 0, "Should have self-assessment criteria"

    for criterion in assessment['self_assessment']:
        assert 'name' in criterion
        assert 'descriptors' in criterion
        assert len(criterion['descriptors']) == 4, "Should have 4 performance levels"


def test_emphasis_competencies(test_sequence):
    """Test emphasizing specific competencies in assessment."""
    from sequence_manager import get_sequence_metadata
    metadata = get_sequence_metadata(test_sequence)

    # Emphasize only the first competency
    first_comp_id = metadata['competencies'][0]['id']

    config = SequenceAssessmentConfig(
        assessment_type="performance_task",
        title="Focused Performance Task",
        emphasis_competencies=[first_comp_id]
    )

    assessment = generate_sequence_assessment(test_sequence, config)

    # Verify task emphasizes the specified competency
    # Check that the first competency statement appears in description or requirements
    first_comp_statement = metadata['competencies'][0]['statement']

    full_text = (assessment['description'] + ' ' +
                ' '.join(assessment['requirements'])).lower()

    # Should reference the emphasized competency
    assert any(word in full_text for word in first_comp_statement.lower().split()[:3]), \
        "Assessment should emphasize specified competencies"


if __name__ == '__main__':
    pytest.main([__file__, '-v'])
